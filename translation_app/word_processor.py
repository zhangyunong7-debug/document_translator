"""
Word document processor for translation
"""
import re
from typing import List, Tuple, Callable, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table


class WordProcessor:
    """Processor for Word documents (.docx)"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = None

    def load(self) -> bool:
        """Load the Word document"""
        try:
            self.document = Document(self.file_path)
            return True
        except Exception as e:
            raise Exception(f"Failed to load Word document: {str(e)}")

    def extract_text_segments(self) -> List[Tuple[str, str, any]]:
        """
        Extract all text segments from the document
        Returns list of tuples: (original_text, element_type, element)
        element_type can be: 'paragraph', 'table_cell', 'header', 'footer'
        """
        segments = []

        if not self.document:
            return segments

        # Extract from paragraphs
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                segments.append((paragraph.text, 'paragraph', paragraph))

        # Extract from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        segments.append((cell.text, 'table_cell', cell))

        # Extract from headers
        for section in self.document.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    segments.append((header.text, 'header', header))

        # Extract from footers
        for section in self.document.sections:
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    segments.append((footer.text, 'footer', footer))

        return segments

    def translate_content(self, translator, source_lang: str, target_lang: str,
                          progress_callback: Optional[Callable] = None) -> bool:
        """
        Translate the entire document content
        """
        if not self.document:
            raise Exception("Document not loaded")

        segments = self.extract_text_segments()
        total = len(segments)

        if total == 0:
            return True

        # Process paragraphs first
        paragraphs_translated = set()
        for text, elem_type, element in segments:
            if elem_type == 'paragraph' and element not in paragraphs_translated:
                try:
                    translated = translator.translate(text, source_lang, target_lang)
                    element.text = translated
                    paragraphs_translated.add(element)
                except Exception as e:
                    print(f"Translation error for paragraph: {e}")

                if progress_callback:
                    progress_callback(len(paragraphs_translated), total)

        # Process tables
        tables_translated = set()
        for text, elem_type, element in segments:
            if elem_type == 'table_cell' and element not in tables_translated:
                try:
                    translated = translator.translate(text, source_lang, target_lang)
                    element.text = translated
                    tables_translated.add(element)
                except Exception as e:
                    print(f"Translation error for table cell: {e}")

                if progress_callback:
                    progress_callback(len(paragraphs_translated) + len(tables_translated), total)

        # Process headers
        for text, elem_type, element in segments:
            if elem_type == 'header':
                try:
                    translated = translator.translate(text, source_lang, target_lang)
                    element.text = translated
                except Exception as e:
                    print(f"Translation error for header: {e}")

        # Process footers
        for text, elem_type, element in segments:
            if elem_type == 'footer':
                try:
                    translated = translator.translate(text, source_lang, target_lang)
                    element.text = translated
                except Exception as e:
                    print(f"Translation error for footer: {e}")

        return True

    def save(self, output_path: str) -> bool:
        """Save the translated document"""
        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            raise Exception(f"Failed to save document: {str(e)}")
